import os
import json
from utils import add_model_table, create_table_with_header
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from response_example import add_response_example
from datetime import datetime

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_description(document, method, method_details, index, sub_index, sub_sub_index):
    document.add_heading(f'{index}.{sub_index}.{sub_sub_index} Description', level=3)
    
    # Descrizioni generiche per i vari tipi di API con sinonimi
    generic_descriptions = {
        'get': [
            'Retrieve {entity}.',
            'Fetch {entity} details.',
            'Get {entity} information.'
        ],
        'get_list': [
            'Retrieve list of {entity_plural}.',
            'Fetch list of {entity_plural}.',
            'Get list of {entity_plural}.'
        ],
        'post': [
            'Create a new {entity} entity.',
            'Add a new {entity}.',
            'Insert a new {entity}.'
        ],
        'delete': [
            'Remove an existing {entity}.',
            'Delete an existing {entity}.',
            'Erase an existing {entity}.'
        ],
        'put': [
            'Update an existing {entity}.',
            'Modify an existing {entity}.'
        ],
        'patch': [
            'Partially update an existing {entity}.',
            'Partially modify an existing {entity}.',
            'Apply partial changes to an existing {entity}.',
            'Make partial updates to an existing {entity}.'
        ]
    }
    
    tags = method_details.get('tags', [])
    entity = tags[0].capitalize() if tags else 'Resource'
    
    entity_plural = entity + 's' if not entity.endswith('s') else entity
    
    is_list = 'list' in method_details.get('operationId', '').lower()
    
    import random
    if method.lower() == 'get' and is_list:
        description = random.choice(generic_descriptions.get('get_list', ['No description provided'])).format(entity=entity, entity_plural=entity_plural)
    else:
        description = random.choice(generic_descriptions.get(method.lower(), ['No description provided'])).format(entity=entity, entity_plural=entity_plural)
    
    document.add_paragraph(description)
    return sub_sub_index + 1

def add_input_parameters(document, parameters, index, sub_index, sub_sub_index):
    if parameters:
        document.add_heading(f'{index}.{sub_index}.{sub_sub_index} Input Parameters', level=3)
        headers = ['Name', 'Description', 'Mandatory']
        table = create_table_with_header(document, headers)
        for parameter in parameters:
            row_cells = table.add_row().cells
            row_cells[0].text = parameter.get('name', 'Name not specified')
            row_cells[1].text = parameter.get('description', 'Description not specified')
            row_cells[2].text = 'Yes' if parameter.get('required', False) else 'No'
        return sub_sub_index + 1
    return sub_sub_index

def add_input_entities_models(document, input_entities, index, sub_index, sub_sub_index):
    for entity_name, entity_details in input_entities.items():
        document.add_heading(f'{index}.{sub_index}.{sub_sub_index} {entity_name} Input Entity Model', level=4)
        sub_sub_index = add_model_table(document, entity_name, entity_details, sub_sub_index)  
    return sub_sub_index

def add_endpoints_description_section(document, swagger_data, index, input_file_path, config):
    try:
        document.add_heading(f'{index}. Endpoints Description', level=1)
        paragraph = document.add_paragraph("BASE_URL description can be found ")
        add_hyperlink(paragraph, config['base_url_description_link'], "here.") 
        sub_index = 1
        for path, path_details in swagger_data.get('paths', {}).items():
            for method, method_details in path_details.items():
                endpoint_name = f'{method.upper()} {path}'
                document.add_heading(f'{index}.{sub_index} {endpoint_name} endpoint', level=2)
                sub_sub_index = 1
                sub_sub_index = add_description(document, method, method_details, index, sub_index, sub_sub_index)
                sub_sub_index = add_input_parameters(document, method_details.get('parameters', []), index, sub_index, sub_sub_index)
                if 'input_entities' in method_details:
                    sub_sub_index = add_input_entities_models(document, method_details['input_entities'], index, sub_index, sub_sub_index)
                elif 'requestBody' in method_details:
                    request_body_content = method_details['requestBody'].get('content', {})
                    for content_type, content_details in request_body_content.items():
                        if 'schema' in content_details:
                            input_entity_name = content_details['schema'].get('$ref', '').split('/')[-1]
                            input_entity_details = swagger_data.get('definitions', {}).get(input_entity_name, {})
                            sub_sub_index = add_input_entities_models(document, {input_entity_name: input_entity_details}, index, sub_index, sub_sub_index)
                sub_sub_index = add_curl_example(document, method, method_details, index, sub_index, sub_sub_index, path, input_file_path, swagger_data)
                sub_sub_index = add_response_example(document, method_details, swagger_data, index, sub_index, sub_sub_index)
                sub_index += 1
        return index + 1
    except Exception as e:
        print(f"Error in add_endpoints_description_section: {e}")
        return index

def add_curl_example(document, method, method_details, index, sub_index, sub_sub_index, path, input_file_path, swagger_data):
    document.add_heading(f'{index}.{sub_index}.{sub_sub_index} cURL', level=3)
    base_url = "{BASE_URL}"
    file_name = os.path.splitext(os.path.basename(input_file_path))[0]
    endpoint = f"{base_url}/{file_name}{path}"
    method_upper = method.upper()
    
    curl_command = f"curl -X '{method_upper}' \\\n  '{endpoint}' \\\n  -H 'accept: application/json'"
    
    query_params = []
    for param in method_details.get('parameters', []):
        if param.get('in') == 'query':
            param_name = param.get('name', 'name_not_specified')
            param_type = param.get('type', 'string')
            if param_type == 'string':
                param_value = 'example_string'
            elif param_type == 'integer':
                param_value = 123
            elif param_type == 'boolean':
                param_value = 'true'
            elif param_type == 'date':
                param_value = datetime.now().strftime('%Y-%m-%d')
            else:
                param_value = 'example_value'
            query_params.append(f"{param_name}={param_value}")
    
    if query_params:
        query_string = '&'.join(query_params)
        curl_command = curl_command.replace(f"'{endpoint}'", f"'{endpoint}?{query_string}'")
    
    if method_upper in ['POST', 'PUT', 'PATCH']:
        request_body = method_details.get('requestBody', {}).get('content', {}).get('application/json', {}).get('example', {})
        if not request_body:
            parameters = method_details.get('parameters', [])
            for param in parameters:
                if param.get('in') == 'body' and '$ref' in param.get('schema', {}):
                    schema_ref = param['schema']['$ref']
                    schema_name = schema_ref.split('/')[-1]
                    if schema_name in swagger_data.get('definitions', {}):
                        request_body = generate_example_from_schema(schema_name, swagger_data.get('definitions', {}))
                        break
            else:
                schema_ref = method_details.get('requestBody', {}).get('content', {}).get('application/json', {}).get('schema', {}).get('$ref', '')
                if schema_ref:
                    schema_name = schema_ref.split('/')[-1]
                    if schema_name in swagger_data.get('definitions', {}):
                        request_body = generate_example_from_schema(schema_name, swagger_data.get('definitions', {}))
                    else:
                        print(f"Schema {schema_name} not found in definitions.")
                else:
                    request_body_schema = method_details.get('requestBody', {}).get('content', {}).get('application/json', {}).get('schema', {})
                    if isinstance(request_body_schema, dict):
                        schema_name = request_body_schema.get('$ref', '').split('/')[-1]
                        if schema_name in swagger_data.get('definitions', {}):
                            request_body = generate_example_from_schema(schema_name, swagger_data.get('definitions', {}))
                        else:
                            print(f"Schema {schema_name} not found in definitions.")
                    else:
                        request_body = {}
        
        print(f"Generated request body: {json.dumps(request_body, indent=2)}")
        
        curl_command += f" \\\n  -H 'Content-Type: application/json' \\\n  -d '{json.dumps(request_body, indent=2)}'"
    
    document.add_paragraph(curl_command)
    return sub_sub_index + 1

def generate_example_from_schema(schema_name, definitions):
    schema = definitions.get(schema_name, {})
    if 'properties' not in schema:
        return {}
    example = {}
    for prop, details in schema['properties'].items():
        if 'type' in details:
            if details['type'] == 'string':
                example[prop] = 'string'
            elif details['type'] == 'integer':
                example[prop] = 123  
            elif details['type'] == 'boolean':
                example[prop] = True
            elif details['type'] == 'array':
                if 'items' in details and 'type' in details['items']:
                    if details['items']['type'] == 'string':
                        example[prop] = ['string']
                    elif details['items']['type'] == 'integer':
                        example[prop] = [123]
                    elif details['items']['type'] == 'boolean':
                        example[prop] = [True]
                    elif details['items']['type'] == 'object':
                        example[prop] = [generate_example_from_schema(details['items'].get('$ref', '').split('/')[-1], definitions)]
                elif 'items' in details and '$ref' in details['items']:
                    ref_name = details['items']['$ref'].split('/')[-1]
                    example[prop] = [generate_example_from_schema(ref_name, definitions)]
            elif details['type'] == 'object':
                example[prop] = generate_example_from_schema(details.get('$ref', '').split('/')[-1], definitions)
        elif '$ref' in details:
            ref_name = details['$ref'].split('/')[-1]
            example[prop] = generate_example_from_schema(ref_name, definitions)
    return example
