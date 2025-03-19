import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
def read_swagger_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            swagger_data = json.load(file)
        return swagger_data
    except Exception as e:
        print(f"Error reading Swagger file: {e}")
        return None

def set_cell_background(cell, color):
    """Set the background color of a cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)

def format_description(attribute):
    """Format the description based on the attribute name."""
    if attribute == 'playerId':
        return 'Player id'
    elif attribute == 'PlayerStatusNow':
        return 'Player status now'
    else:
        words = []
        current_word = ""
        for char in attribute:
            if char.isupper() and current_word:
                words.append(current_word)
                current_word = char
            else:
                current_word += char
        words.append(current_word)
        return ' '.join(words).capitalize()

def format_type(attribute, type_value):
    """Format the type based on the attribute name and type value."""
    if type_value == 'Not specified':
        return f'{attribute[0].upper()}{attribute[1:]} entity'
    return type_value

def create_table_with_header(document, headers):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], 'f3f3f3')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    return table

def add_model_table(document, model_name, model, section_index, model_index):
    document.add_heading(f'{section_index}.{model_index} {model_name} entity', level=3)
    headers = ['Name', 'Description', 'Type']
    table = create_table_with_header(document, headers)
    
    for attribute, details in model.get('properties', {}).items():
        row_cells = table.add_row().cells
        row_cells[0].text = attribute
        row_cells[1].text = format_description(attribute)
        row_cells[2].text = format_type(attribute, details.get('type', 'Not specified'))
    return model_index + 1
